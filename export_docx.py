#!/usr/bin/env python3
"""
Module: DOCX (Word) export for AI Document Writer
Version: 1.0.0
Development Iteration: v1

Project: AI Document Writer
Developer: Kent Benson
Created: 2026-01-29

Enhancement: Initial implementation

Features:
- Convert plain text documents to Word (.docx) format
- Auto-detect headings, bullets, numbered lists
- Professional formatting

UV ENVIRONMENT: Run with `uv run python export_docx.py`

INSTALLATION:
uv add python-docx
"""

import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import DRAFTS_DIR

logger = logging.getLogger(__name__)


def export_to_docx(
    text: str,
    title: str = "Document",
    output_path: Optional[str] = None,
) -> Optional[str]:
    """
    Export plain text document to Word (.docx) format.

    Args:
        text: The document text to export
        title: Title for the document header
        output_path: Where to save. If None, saves to drafts directory.

    Returns:
        Path to saved DOCX, or None on error
    """
    if not text.strip():
        logger.error("No text to export")
        return None

    try:
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add date
        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in date_para.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.size = Pt(10)

        doc.add_paragraph("")  # Spacer

        # Process text line by line
        lines = text.split('\n')
        for line in lines:
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                continue

            # ALL CAPS = section heading
            if stripped.isupper() and len(stripped) > 3 and len(stripped) < 80:
                doc.add_heading(stripped.title(), level=1)

            # Lines ending with colon = sub-heading
            elif stripped.endswith(':') and len(stripped) < 60:
                doc.add_heading(stripped, level=2)

            # Bullet points
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')

            # Numbered items
            elif re.match(r'^\d+[\.\)]\s', stripped):
                clean = re.sub(r'^\d+[\.\)]\s', '', stripped)
                doc.add_paragraph(clean, style='List Number')

            # Regular paragraph
            else:
                doc.add_paragraph(stripped)

        # Determine output path
        if output_path is None:
            clean_name = re.sub(r'[^\w\s-]', '', title)[:30].replace(' ', '_')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{clean_name}_{timestamp}.docx"
            output_path = str(DRAFTS_DIR / filename)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        doc.save(output_path)
        logger.info(f"DOCX exported: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Failed to export DOCX: {e}", exc_info=True)
        return None
