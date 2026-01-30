"""Tests for export_docx.py — DOCX generation with real python-docx."""

from pathlib import Path

import pytest
from docx import Document

from export_docx import export_to_docx


class TestExportToDocx:
    """Tests for export_to_docx()."""

    def test_creates_docx_file(self, tmp_drafts_dir, sample_document_text):
        path = export_to_docx(sample_document_text, title="Test Doc")
        assert path is not None
        assert Path(path).exists()
        assert path.endswith(".docx")

    def test_returns_none_for_empty_text(self, tmp_drafts_dir):
        assert export_to_docx("") is None
        assert export_to_docx("   ") is None
        assert export_to_docx("\n\n") is None

    def test_uses_custom_output_path(self, tmp_drafts_dir, sample_document_text):
        custom = str(tmp_drafts_dir / "custom_output.docx")
        path = export_to_docx(sample_document_text, title="Custom", output_path=custom)
        assert path == custom
        assert Path(custom).exists()

    def test_auto_generates_filename_with_sanitized_title(self, tmp_drafts_dir):
        path = export_to_docx("Some content here.", title="My Report!")
        filename = Path(path).name
        assert "!" not in filename
        assert filename.endswith(".docx")

    def test_title_sanitization_truncates_long_names(self, tmp_drafts_dir):
        long_title = "A" * 100
        path = export_to_docx("Content.", title=long_title)
        filename = Path(path).name
        title_part = filename.split("_2")[0]  # Before timestamp
        assert len(title_part) <= 30

    def test_contains_title_heading(self, tmp_drafts_dir, sample_document_text):
        path = export_to_docx(sample_document_text, title="My Report Title")
        doc = Document(path)
        # First paragraph should be a heading with the title
        first = doc.paragraphs[0]
        assert first.text == "My Report Title"
        assert first.style.name in ("Title", "Heading 0") or first.style.name.startswith("Heading")

    def test_all_caps_become_headings(self, tmp_drafts_dir):
        text = "EXECUTIVE SUMMARY\n\nThis is the summary.\n"
        path = export_to_docx(text, title="Test")
        doc = Document(path)
        # Find the ALL CAPS line — it should be a heading (title-cased)
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Executive Summary" in heading_texts

    def test_bullets_preserved(self, tmp_drafts_dir):
        text = "Items:\n- First item\n- Second item\n* Third item\n"
        path = export_to_docx(text, title="Test")
        doc = Document(path)
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullet_paras) == 3
        assert bullet_paras[0].text == "First item"

    def test_numbered_items_preserved(self, tmp_drafts_dir):
        text = "Steps:\n1. Do this\n2) Do that\n"
        path = export_to_docx(text, title="Test")
        doc = Document(path)
        numbered_paras = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered_paras) == 2

    def test_file_is_valid_docx(self, tmp_drafts_dir, sample_document_text):
        path = export_to_docx(sample_document_text, title="Valid DOCX")
        # Should be re-openable without error
        doc = Document(path)
        assert len(doc.paragraphs) > 0

    def test_creates_parent_directories(self, tmp_path, sample_document_text):
        nested = tmp_path / "a" / "b" / "c" / "test.docx"
        path = export_to_docx(sample_document_text, output_path=str(nested))
        assert Path(path).exists()
